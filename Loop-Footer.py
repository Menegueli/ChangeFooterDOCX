# A ideia é fazer um loop no qual eu copio o mesmo arquivo x vezes
# Apenas mudando o código do rodapé

from docx import Document  #Importo a biblioteca docx após fazer um pip3 install python-docx
document = Document('Documento-Gabriel1.docx')  #Seleciono o arquivo que quero copiar
section = document.sections[0]  #Seleciono a parte para alterar (Normalmente wdocx tem apenas uma seção, logo, 0)
footer = section.footer  #Na Seção escolhida eu defino qual parte do layout quero mexer, no caso o footer (rodapé)

for x in range(2,151):  #Faço um loop com a quantidade de vezes que quero repetir, lembrando que 1 é o documento original, então começo no 2
    i=x-1 #Faço uma variavel no qual ele vai pegar o valor do arquivo antigo para substituir pelo novo

    footer.paragraphs[1].text  = footer.paragraphs[1].text.replace(f'{i}', f'{x}') #Aqui fazemos o replace do rodapé pegando o valor antigo (i) e dando replace para o novo (x)


    document.save(f'Documento-Gabriel{x}.docx') #Salvando o documento declarando a versão com qual número é com o (x)